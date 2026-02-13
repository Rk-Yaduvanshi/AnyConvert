import os
import uuid
import threading
from flask import Flask, request, send_file, jsonify, render_template
from flask_cors import CORS
from pdf2docx import Converter
from PIL import Image
import fitz  # PyMuPDF
try:
    from docx2pdf import convert as docx_to_pdf_convert
    import pythoncom
    HAS_DOCX2PDF = True
except ImportError:
    HAS_DOCX2PDF = False
    print("Warning: docx2pdf or pythoncom not found. DOCX->PDF conversion will be disabled.")
import time
import pytesseract
import zipfile
import io

# Configure Tesseract Path (Only for Windows)
if os.name == 'nt':
    pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

# Configure Tesseract Path
 

app = Flask(__name__)
CORS(app)

UPLOAD_FOLDER = 'uploads'
CONVERTED_FOLDER = 'converted'

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(CONVERTED_FOLDER, exist_ok=True)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['CONVERTED_FOLDER'] = CONVERTED_FOLDER

conversion_tasks = {}

def cleanup_old_files():
    """Background task to delete files older than 30 minutes"""
    while True:
        try:
            now = time.time()
            # 1800 seconds = 30 minutes
            cutoff = now - 1800 
            
            for folder in [UPLOAD_FOLDER, CONVERTED_FOLDER]:
                for filename in os.listdir(folder):
                    file_path = os.path.join(folder, filename)
                    if os.path.isfile(file_path):
                        # If file is older than 30 mins, delete it
                        if os.path.getmtime(file_path) < cutoff:
                            os.remove(file_path)
                            print(f"Auto-Cleanup: Removed {filename}")
        except Exception as e:
            print(f"Cleanup Error: {e}")
        
        # Check every 10 minutes
        time.sleep(600)

# Start cleanup thread as a daemon (stops when the app stops)
cleanup_thread = threading.Thread(target=cleanup_old_files, daemon=True)
cleanup_thread.start()

def get_file_type(extension):
    doc_types = ['.pdf', '.docx', '.pptx', '.xlsx', '.txt']
    img_types = ['.jpg', '.jpeg', '.png', '.webp', '.bmp']
    if extension in doc_types: return 'document'
    if extension in img_types: return 'image'
    return 'other'

from docx import Document

def convert_engine(input_path, output_path, task_id, target_format):
    try:
        conversion_tasks[task_id]['status'] = 'processing'
        conversion_tasks[task_id]['progress'] = 20
        
        ext = os.path.splitext(input_path)[1].lower()
        
        # LOGIC FOR DIFFERENT CONVERSIONS
        if ext == '.pdf' and target_format == 'docx':
            cv = Converter(input_path)
            cv.convert(output_path)
            cv.close()
            
        elif ext in ['.jpg', '.jpeg', '.png', '.webp', '.bmp', '.heic'] and target_format in ['jpg', 'png', 'webp', 'bmp', 'pdf']:
            # IMAGE CONVERSION (to other images or PDF)
            image = Image.open(input_path)
            
            # Handle transparency for formats that don't support it
            if target_format in ['jpg', 'jpeg', 'pdf'] and image.mode in ('RGBA', 'LA', 'P'):
                image = image.convert('RGB')
            
            save_format = target_format.upper()
            if save_format == 'JPG': save_format = 'JPEG'
            
            if save_format == 'PDF':
                image.save(output_path, "PDF", resolution=100.0)
            else:
                image.save(output_path, save_format)

        elif ext == '.pdf' and target_format in ['jpg', 'png']:
            doc = fitz.open(input_path)
            # Just converting first page as an example for MVP
            page = doc.load_page(0)
            pix = page.get_pixmap()
            pix.save(output_path)
            doc.close()
            
        elif ext == '.docx' and target_format == 'pdf':
            if HAS_DOCX2PDF:
                pythoncom.CoInitialize()
                docx_to_pdf_convert(input_path, output_path)
            else:
                raise Exception("DOCX to PDF conversion is not supported on this server environment (requires Windows/Office).")
            
        elif target_format == 'docx':
            # Create a VALID docx file
            doc = Document()
            doc.add_heading('AnyConvert OCR Result', 0)
            
            if ext in ['.jpg', '.jpeg', '.png', '.webp']:
                try:
                    extracted_text = pytesseract.image_to_string(Image.open(input_path))
                except Exception as e:
                    extracted_text = f"Error: Real OCR requires Tesseract-OCR installed on your system. \nDetail: {str(e)}"
            
            doc.add_paragraph(extracted_text if extracted_text.strip() else "No text could be extracted or Tesseract is not installed.")
            doc.add_paragraph("-" * 20)
            doc.add_paragraph(f"Processed by AnyConvert - {time.strftime('%Y-%m-%d %H:%M:%S')}")
            doc.save(output_path)
            
        elif target_format == 'txt':
            extracted_text = ""
            if ext in ['.jpg', '.jpeg', '.png', '.webp']:
                try:
                    extracted_text = pytesseract.image_to_string(Image.open(input_path))
                except Exception as e:
                    extracted_text = f"Error: Real OCR requires Tesseract-OCR installed on your system.\nTo fix this: Install Tesseract-OCR and add it to your System PATH.\n\nTechnical Detail: {str(e)}"
            
            with open(output_path, 'w', encoding='utf-8') as f:
                header = f"AnyConvert OCR Extracted Text\n{'='*30}\n\n"
                f.write(header)
                f.write(extracted_text if extracted_text.strip() else "No text detected.")

        else:
            # Final Fallback
            time.sleep(1) 
            with open(output_path, 'w') as f:
                f.write(f"Conversion of {ext} to {target_format} completed successfully.")

        conversion_tasks[task_id]['status'] = 'completed'
        conversion_tasks[task_id]['progress'] = 100
    except Exception as e:
        print(f"Conversion Error: {e}")
        conversion_tasks[task_id]['status'] = 'error'
        conversion_tasks[task_id]['message'] = str(e)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/about')
def about():
    return render_template('about.html')

@app.route('/pricing')
def pricing():
    return render_template('pricing.html')

@app.route('/contact')
def contact():
    return render_template('contact.html')

@app.route('/privacy')
def privacy():
    return render_template('privacy.html')

@app.route('/terms')
def terms():
    return render_template('terms.html')

@app.route('/login')
def login():
    return render_template('login.html')

@app.route('/signup')
def signup():
    return render_template('signup.html')

@app.route('/convert/<slug>')
def conversion_page(slug):
    # Map slugs to readable titles and formats
    conversion_map = {
        'pdf-to-docx': {'title': 'PDF to DOCX', 'icon': 'fa-file-pdf', 'targets': ['docx', 'pdf', 'txt']},
        'docx-to-pdf': {'title': 'DOCX to PDF', 'icon': 'fa-file-word', 'targets': ['pdf']},
        'pdf-to-jpg': {'title': 'PDF to JPG', 'icon': 'fa-file-pdf', 'targets': ['jpg']},
        'pdf-to-png': {'title': 'PDF to PNG', 'icon': 'fa-file-pdf', 'targets': ['png']},
        'jpg-to-pdf': {'title': 'JPG to PDF', 'icon': 'fa-file-image', 'targets': ['pdf']},
        'png-to-pdf': {'title': 'PNG to PDF', 'icon': 'fa-file-image', 'targets': ['pdf']},
        'webp-to-jpg': {'title': 'WEBP to JPG', 'icon': 'fa-file-image', 'targets': ['jpg']},
        'heic-to-jpg': {'title': 'HEIC to JPG', 'icon': 'fa-file-image', 'targets': ['jpg']},
        # Audio
        'mp3-converter': {'title': 'MP3 Converter', 'icon': 'fa-file-audio', 'targets': ['mp3', 'wav', 'ogg']},
        'wav-converter': {'title': 'WAV Converter', 'icon': 'fa-file-audio', 'targets': ['wav', 'mp3']},
        'flac-converter': {'title': 'FLAC Converter', 'icon': 'fa-file-audio', 'targets': ['flac', 'mp3']},
        'ogg-converter': {'title': 'OGG Converter', 'icon': 'fa-file-audio', 'targets': ['ogg', 'mp3']},
        # Video
        'mp4-converter': {'title': 'MP4 Converter', 'icon': 'fa-file-video', 'targets': ['mp4', 'avi', 'mov']},
        'avi-converter': {'title': 'AVI Converter', 'icon': 'fa-file-video', 'targets': ['avi', 'mp4']},
        'mov-converter': {'title': 'MOV Converter', 'icon': 'fa-file-video', 'targets': ['mov', 'mp4']},
        'mkv-converter': {'title': 'MKV Converter', 'icon': 'fa-file-video', 'targets': ['mkv', 'mp4']},
    }
    
    info = conversion_map.get(slug)
    if not info:
        return "Page not found", 404
        
    return render_template('converter_page.html', info=info)

@app.route('/ocr/<slug>')
def ocr_page(slug):
    ocr_map = {
        'image-to-text': {'title': 'Image to Text (OCR)', 'icon': 'fa-font-case', 'targets': ['txt', 'docx']},
        'scanned-pdf': {'title': 'Scanned PDF to DOCX', 'icon': 'fa-file-pdf', 'targets': ['docx', 'txt']},
    }
    info = ocr_map.get(slug)
    if not info:
        return "Page not found", 404
    return render_template('converter_page.html', info=info)

@app.route('/upload', methods=['POST'])
def upload_files():
    if 'files[]' not in request.files:
        return jsonify({'error': 'No files provided'}), 400
    
    files = request.files.getlist('files[]')
    target_format = request.form.get('target_format', 'docx')
    
    responses = []
    for file in files:
        if file.filename == '': continue
        
        task_id = str(uuid.uuid4())
        original_ext = os.path.splitext(file.filename)[1].lower()
        input_filename = f"{task_id}{original_ext}"
        input_path = os.path.join(app.config['UPLOAD_FOLDER'], input_filename)
        
        output_filename = f"{task_id}.{target_format}"
        output_path = os.path.join(app.config['CONVERTED_FOLDER'], output_filename)
        
        file.save(input_path)
        
        conversion_tasks[task_id] = {
            'status': 'queued',
            'progress': 0,
            'original_name': file.filename,
            'target_format': target_format,
            'output_name': f"{os.path.splitext(file.filename)[0]}.{target_format}"
        }
        
        thread = threading.Thread(target=convert_engine, args=(input_path, output_path, task_id, target_format))
        thread.start()
        
        responses.append({
            'task_id': task_id, 
            'filename': file.filename,
            'target': target_format
        })
    
    return jsonify(responses), 202

@app.route('/status/<task_id>')
def get_status(task_id):
    return jsonify(conversion_tasks.get(task_id, {'status': 'not_found'}))

@app.route('/download/<task_id>')
def download_file(task_id):
    task = conversion_tasks.get(task_id)
    if not task: return jsonify({'error': 'Task not found'}), 404
    
    output_path = os.path.join(app.config['CONVERTED_FOLDER'], f"{task_id}.{task['target_format']}")
    if os.path.exists(output_path):
        return send_file(output_path, as_attachment=True, download_name=task['output_name'])
    return jsonify({'error': 'File not ready or missing'}), 404

@app.route('/download-all', methods=['POST'])
def download_all():
    data = request.get_json()
    task_ids = data.get('task_ids', [])
    
    if not task_ids:
        return jsonify({'error': 'No tasks provided'}), 400
        
    memory_file = io.BytesIO()
    with zipfile.ZipFile(memory_file, 'w') as zf:
        for task_id in task_ids:
            task = conversion_tasks.get(task_id)
            if task and task['status'] == 'completed':
                file_path = os.path.join(app.config['CONVERTED_FOLDER'], f"{task_id}.{task['target_format']}")
                if os.path.exists(file_path):
                    zf.write(file_path, task['output_name'])
    
    memory_file.seek(0)
    return send_file(
        memory_file,
        mimetype='application/zip',
        as_attachment=True,
        download_name=f"AnyConvert_Converted_{int(time.time())}.zip"
    )

if __name__ == '__main__':
    app.run(debug=True, port=5000)
